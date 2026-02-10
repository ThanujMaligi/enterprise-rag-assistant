import os
import json
import csv
import hashlib
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

@dataclass
class Document:
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: Optional[str] = None

    def __post_init__(self):
        if not self.doc_id:
            hash_str = f"{self.metadata.get('source', '')}_{self.metadata.get('page', 0)}_{self.content[:100]}"
            self.doc_id = hashlib.sha256(hash_str.encode('utf-8')).hexdigest()[:16]

class EnterpriseDocumentLoader:
    """Multi-format enterprise document loader supporting PDF, DOCX, TXT, MD, CSV, JSON."""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.csv', '.json'}

    @classmethod
    def load_file(cls, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}. Supported formats: {cls.SUPPORTED_EXTENSIONS}")

        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        if ext == '.pdf':
            return cls._load_pdf(file_path, filename, file_size)
        elif ext == '.docx':
            return cls._load_docx(file_path, filename, file_size)
        elif ext in ('.txt', '.md'):
            return cls._load_text(file_path, filename, file_size)
        elif ext == '.csv':
            return cls._load_csv(file_path, filename, file_size)
        elif ext == '.json':
            return cls._load_json(file_path, filename, file_size)
        return []

    @classmethod
    def _load_pdf(cls, file_path: str, filename: str, file_size: int) -> List[Document]:
        documents = []
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    documents.append(Document(
                        content=text.strip(),
                        metadata={
                            "source": filename,
                            "file_path": file_path,
                            "file_type": "pdf",
                            "file_size": file_size,
                            "page": idx + 1,
                            "total_pages": len(reader.pages)
                        }
                    ))
        except Exception as e:
            # Fallback text reading if pypdf fails
            documents.append(Document(
                content=f"[Error reading PDF pages: {str(e)}]",
                metadata={"source": filename, "file_path": file_path, "file_type": "pdf", "error": str(e)}
            ))
        return documents

    @classmethod
    def _load_docx(cls, file_path: str, filename: str, file_size: int) -> List[Document]:
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            return [Document(
                content=full_text,
                metadata={
                    "source": filename,
                    "file_path": file_path,
                    "file_type": "docx",
                    "file_size": file_size,
                    "paragraph_count": len(paragraphs)
                }
            )]
        except Exception as e:
            return [Document(
                content=f"[Error reading DOCX: {str(e)}]",
                metadata={"source": filename, "file_path": file_path, "file_type": "docx", "error": str(e)}
            )]

    @classmethod
    def _load_text(cls, file_path: str, filename: str, file_size: int) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read().strip()
        ext = os.path.splitext(file_path)[1].lower()[1:]
        return [Document(
            content=text,
            metadata={
                "source": filename,
                "file_path": file_path,
                "file_type": ext,
                "file_size": file_size,
                "character_count": len(text)
            }
        )]

    @classmethod
    def _load_csv(cls, file_path: str, filename: str, file_size: int) -> List[Document]:
        documents = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.DictReader(f)
            for idx, row in enumerate(reader):
                row_str = ", ".join([f"{k}: {v}" for k, v in row.items() if v])
                if row_str:
                    documents.append(Document(
                        content=row_str,
                        metadata={
                            "source": filename,
                            "file_path": file_path,
                            "file_type": "csv",
                            "file_size": file_size,
                            "row_index": idx + 1
                        }
                    ))
        return documents

    @classmethod
    def _load_json(cls, file_path: str, filename: str, file_size: int) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            data = json.load(f)
        if isinstance(data, list):
            content = "\n".join([json.dumps(item, indent=2) for item in data])
        else:
            content = json.dumps(data, indent=2)
        return [Document(
            content=content,
            metadata={
                "source": filename,
                "file_path": file_path,
                "file_type": "json",
                "file_size": file_size
            }
        )]
